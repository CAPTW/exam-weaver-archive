# src/exporter/docx.py
"""Word 문서(DOCX) 내보내기"""

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import logging
import random
import json
import io
import re

from ..parser.formatting import normalize_private_math_glyphs, repair_extracted_text_artifacts
from ..parser.patterns import NUMBER_TO_CHOICE_SYMBOL, CHOICE_SYMBOL_TO_NUMBER
from ..parser.question import ALL_CHOICES_CORRECT
from ..parser.aligned_choice_table import is_aligned_choice_format
from ..parser.table_format import (
    effective_table_render_mode,
    parse_format_payload,
    resolve_table_anchor,
)
from ..choice_markers import (
    DEFAULT_CHOICE_MARKER_STYLE,
    choice_marker,
    normalize_choice_marker_style,
)
from .table_layout import fallback_table_layout, resolve_table_layout

logger = logging.getLogger(__name__)

class DocxExporter:
    FONT_NAME = '경기천년제목OTF Light'
    BODY_LINE_TWIPS = '160'
    COLUMN_SPACE_TWIPS = '425'

    def __init__(
        self,
        choice_marker_style=DEFAULT_CHOICE_MARKER_STYLE,
        table_render_mode="auto",
    ):
        self.choice_marker_style = normalize_choice_marker_style(choice_marker_style)
        self.table_render_mode = self._normalize_table_render_mode(table_render_mode)
        self.warnings = []

    def set_choice_marker_style(self, style):
        self.choice_marker_style = normalize_choice_marker_style(style)

    def set_table_render_mode(self, mode):
        self.table_render_mode = self._normalize_table_render_mode(mode)

    @staticmethod
    def _normalize_table_render_mode(mode):
        mode = str(mode or "auto").lower()
        return mode if mode in {"auto", "image", "native"} else "auto"

    def export(
        self,
        title: str,
        questions: list,
        output_path: str,
        shuffle_choices: bool = False,
        include_answer_key: bool = False,
        sections: list = None
    ):
        """
        Export questions to a DOCX document matching the exam handout style.

        Args:
            title: Document title. Use newlines for title/subtitle lines.
            questions: Questions to export.
            output_path: Destination DOCX path.
            shuffle_choices: Shuffle four-choice options before rendering.
            include_answer_key: Append a separate answer key page when needed.
            sections: Optional grouped sections, each with title and questions.
        """
        doc = Document()
        self.warnings = []

        self._set_document_defaults(doc)

        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)

        self._set_columns(section, 2)

        for title_line in self._split_title(title):
            header = doc.add_paragraph()
            self._format_paragraph(header, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = header.add_run(f"{title_line} ")
            self._format_run(run, size_pt=16, bold=True)

        rng = random.Random()
        answer_key = []
        display_number = 1
        if sections:
            for section_spec in sections:
                last_group_id = None
                section_title = (section_spec or {}).get('title')
                if section_title:
                    self._add_subject_heading(doc, section_title)
                    spacer = doc.add_paragraph()
                    self._format_paragraph(spacer)
                for q in (section_spec or {}).get('questions') or []:
                    last_group_id = self._add_group_shared_passage_if_needed(
                        doc,
                        q,
                        last_group_id,
                    )
                    answer = self._add_question(
                        doc,
                        q,
                        display_number=display_number,
                        shuffle_choices=shuffle_choices,
                        rng=rng
                    )
                    answer_key.append((display_number, answer))
                    display_number += 1
                    spacer = doc.add_paragraph()
                    self._format_paragraph(spacer)
        else:
            last_group_id = None
            for q in questions:
                last_group_id = self._add_group_shared_passage_if_needed(
                    doc,
                    q,
                    last_group_id,
                )
                answer = self._add_question(
                    doc,
                    q,
                    display_number=display_number,
                    shuffle_choices=shuffle_choices,
                    rng=rng
                )
                answer_key.append((display_number, answer))
                display_number += 1
                spacer = doc.add_paragraph()
                self._format_paragraph(spacer)

        if include_answer_key and answer_key:
            self._add_answer_key(doc, answer_key)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"Exported to {output_path}")

    def _add_group_shared_passage_if_needed(self, doc, question, last_group_id):
        group_id = question.get('group_id')
        shared_passage = (
            question.get('shared_passage')
            or question.get('group_shared_text')
        )
        if group_id is not None and shared_passage and group_id != last_group_id:
            self._add_shared_passage(doc, shared_passage)
        return group_id if group_id is not None else None

    def _add_shared_passage(self, doc, shared_passage):
        paragraph = doc.add_paragraph()
        self._format_paragraph(paragraph)
        label = paragraph.add_run("[공통지문] ")
        self._format_run(label, bold=True)
        self._add_text_with_inline_math(
            paragraph,
            self._normalize_plain_export_text(shared_passage),
        )

    def _add_subject_heading(self, doc, title):
        paragraph = doc.add_paragraph()
        self._format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = paragraph.add_run(str(title or '').strip())
        self._format_run(run, size_pt=16, bold=True)

    def _add_question(
        self,
        doc,
        q,
        display_number: int,
        shuffle_choices: bool = False,
        rng: random.Random = None
    ):
        """Add a question and its choices. Returns answer symbol or None."""
        p = self._add_text_with_format_tables(
            doc,
            q.get('question_text', ''),
            q.get('question_format_json'),
            prefix=f"{display_number}. "
        )
        question_image_path = q.get('image_path')
        if self._image_exists(question_image_path):
            self._keep_paragraph_together(p, keep_with_next=True)
        self._add_image(
            doc,
            question_image_path,
            width_mm=60,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=bool(q.get('choices')),
            keep_together=True,
        )

        if q.get('question_type') == 'descriptive':
            model_answer = str(q.get('model_answer') or '').strip()
            if model_answer:
                p = doc.add_paragraph()
                self._format_paragraph(p)
                self._add_formatted_text(
                    p,
                    model_answer,
                    prefix="모범답안: ",
                    size_pt=10,
                )
            return "서술형"

        answer_number = q.get('correct_answer')

        # Choices
        if q.get('choices'):
            choices = [self._normalize_choice(c) for c in q['choices']]

            if shuffle_choices and len(choices) == 4:
                rng = rng or random.Random()
                rng.shuffle(choices)

            if (
                answer_number is not None
                and answer_number != ALL_CHOICES_CORRECT
                and shuffle_choices
                and len(choices) == 4
            ):
                new_answer = None
                for idx, choice in enumerate(choices, start=1):
                    if choice.get('_orig_number') == answer_number:
                        new_answer = idx
                        break
                answer_number = new_answer

            if shuffle_choices and len(choices) == 4:
                for idx, choice in enumerate(choices, start=1):
                    choice['choice_number'] = idx
                    choice['choice_symbol'] = NUMBER_TO_CHOICE_SYMBOL.get(idx, str(idx))

            for choice in choices:
                stored_symbol = choice.get('choice_symbol') or ''
                symbol = choice_marker(
                    choice.get('choice_number'),
                    self.choice_marker_style,
                    fallback=stored_symbol,
                )
                is_correct = (
                    answer_number == ALL_CHOICES_CORRECT
                    or choice.get('choice_number') == answer_number
                )
                choice_format_json = (
                    choice.get('choice_format_json') or choice.get('format_json')
                )
                choice_text = (
                    ''
                    if is_aligned_choice_format(choice_format_json)
                    else choice.get('choice_text', '')
                )
                p = self._add_text_with_format_tables(
                    doc,
                    choice_text,
                    choice_format_json,
                    prefix=f"{symbol} " if symbol else '',
                    size_pt=10,
                    highlight=is_correct
                )
                self._add_image(
                    doc,
                    choice.get('choice_image_path'),
                    width_mm=42,
                    paragraph=p,
                )

        if answer_number is None:
            return None

        if answer_number == ALL_CHOICES_CORRECT:
            return "전원 정답"

        return choice_marker(
            answer_number,
            self.choice_marker_style,
            fallback=NUMBER_TO_CHOICE_SYMBOL.get(answer_number, str(answer_number)),
        )

    def _normalize_choice(self, choice):
        """Normalize choice object to a dict."""
        if isinstance(choice, dict):
            data = dict(choice)
            orig_number = data.get('choice_number')
            if orig_number is None and data.get('choice_symbol'):
                orig_number = CHOICE_SYMBOL_TO_NUMBER.get(data['choice_symbol'])
            data['_orig_number'] = orig_number
            data['choice_image_path'] = data.get('choice_image_path') or data.get('image_path')
            data['choice_format_json'] = data.get('choice_format_json') or data.get('format_json')
            return data

        orig_number = getattr(choice, 'number', None)
        if orig_number is None:
            sym = getattr(choice, 'symbol', None)
            if sym:
                orig_number = CHOICE_SYMBOL_TO_NUMBER.get(sym)

        return {
            'choice_number': getattr(choice, 'number', None),
            'choice_symbol': getattr(choice, 'symbol', None),
            'choice_text': getattr(choice, 'text', None),
            'choice_image_path': getattr(choice, 'choice_image_path', None) or getattr(choice, 'image_path', None),
            'choice_format_json': getattr(choice, 'choice_format_json', None) or getattr(choice, 'format_json', None),
            '_orig_number': orig_number
        }

    def _parse_format_json(self, value):
        if not value:
            return {}
        if isinstance(value, dict):
            return value
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, dict) else {}
        except (TypeError, ValueError):
            return {}

    def _add_formatted_text(
        self,
        paragraph,
        text,
        format_json=None,
        prefix='',
        size_pt=None,
        highlight=False
    ):
        text = str(text or '')
        format_spec = self._parse_format_json(format_json)
        spans = self._valid_spans(text, format_spec.get('spans') or [])
        if not spans:
            text = self._normalize_plain_export_text(text)
            if prefix:
                run = paragraph.add_run(prefix)
                self._format_run(run, size_pt=size_pt, highlight=highlight)
            self._add_text_with_inline_math(
                paragraph,
                text,
                size_pt=size_pt,
                highlight=highlight
            )
            return

        if prefix:
            run = paragraph.add_run(prefix)
            self._format_run(run, size_pt=size_pt, highlight=highlight)

        position = 0
        for start, end, options in spans:
            if start > position:
                self._add_text_with_inline_math(
                    paragraph,
                    self._normalize_plain_export_text(text[position:start], strip=False),
                    size_pt=size_pt,
                    highlight=highlight
                )
            if options.get('latex'):
                self._add_math_from_latex(paragraph, options.get('latex'))
            else:
                run = paragraph.add_run(text[start:end])
                self._format_run(
                    run,
                    size_pt=size_pt,
                    highlight=highlight,
                    underline=bool(options.get('underline'))
                )
            position = end

        if position < len(text) or not spans:
            self._add_text_with_inline_math(
                paragraph,
                self._normalize_plain_export_text(text[position:], strip=False),
                size_pt=size_pt,
                highlight=highlight
            )

    def _normalize_plain_export_text(self, text, strip=True):
        """Remove PDF extraction layout artifacts before writing plain DOCX text."""
        if strip:
            return repair_extracted_text_artifacts(text)

        raw = str(text or '')
        leading = re.match(r'^\s*', raw).group(0)
        trailing = re.search(r'\s*$', raw).group(0)
        core_end = len(raw) - len(trailing) if trailing else len(raw)
        core = raw[len(leading):core_end]
        return leading + repair_extracted_text_artifacts(core) + trailing

    def _normalize_private_math_glyphs(self, text):
        return normalize_private_math_glyphs(text)

    def _private_digit(self, value):
        return {
            '\ue034': '1',
            '\ue035': '2',
            '\ue036': '3',
            '\ue037': '4',
        }.get(value, value)

    def _add_text_with_inline_math(self, paragraph, text, size_pt=None, highlight=False):
        position = 0
        for match in re.finditer(r'√\s*([0-9]+|[A-Za-z]+)', str(text or '')):
            if match.start() > position:
                run = paragraph.add_run(text[position:match.start()])
                self._format_run(run, size_pt=size_pt, highlight=highlight)
            self._add_radical_math(paragraph, match.group(1))
            position = match.end()

        if position < len(text) or not text:
            run = paragraph.add_run(text[position:])
            self._format_run(run, size_pt=size_pt, highlight=highlight)

    def _add_math_from_latex(self, paragraph, latex):
        latex = str(latex or '').strip()
        element = self._math_element_from_latex(latex)
        if element is None:
            element = self._math_text_element(latex)
        paragraph._p.append(element)

    def _math_element_from_latex(self, latex):
        sqrt_match = re.fullmatch(r'\\?sqrt\{(.+)\}', latex)
        if sqrt_match:
            return self._radical_math_element(sqrt_match.group(1))

        overline_match = re.fullmatch(r'\\?overline\{(.+)\}', latex)
        if overline_match:
            return self._bar_math_element(overline_match.group(1))

        frac_match = re.fullmatch(r'\\?frac\{(.+)\}\{(.+)\}', latex)
        if frac_match:
            return self._fraction_math_element(frac_match.group(1), frac_match.group(2))

        if self._looks_like_linear_latex_formula(latex):
            return self._linear_math_element(latex)

        return None

    def _add_radical_math(self, paragraph, inner):
        paragraph._p.append(self._radical_math_element(inner))

    def _math_text_element(self, text):
        math = OxmlElement('m:oMath')
        math.append(self._math_run(text))
        return math

    def _radical_math_element(self, inner):
        math = OxmlElement('m:oMath')
        math.append(self._radical_component(inner))
        return math

    def _radical_component(self, inner):
        rad = OxmlElement('m:rad')
        rad_pr = OxmlElement('m:radPr')
        deg_hide = OxmlElement('m:degHide')
        deg_hide.set(qn('m:val'), '1')
        rad_pr.append(deg_hide)
        rad.append(rad_pr)
        rad.append(OxmlElement('m:deg'))
        expr = OxmlElement('m:e')
        expr.append(self._math_run(inner))
        rad.append(expr)
        return rad

    def _looks_like_linear_latex_formula(self, latex):
        return bool(re.search(r'=|\\sqrt|\\times|\\cdot|\\theta|\\cos|\\sin|\\tan', str(latex or '')))

    def _linear_math_element(self, latex):
        math = OxmlElement('m:oMath')
        position = 0
        text = str(latex or '')
        for match in re.finditer(r'\\sqrt\{([^{}]+)\}', text):
            if match.start() > position:
                self._append_math_text(math, text[position:match.start()])
            math.append(self._radical_component(match.group(1)))
            position = match.end()

        if position < len(text):
            self._append_math_text(math, text[position:])
        return math

    def _append_math_text(self, math, latex_text):
        display = self._latex_text_to_display_text(latex_text)
        if display:
            math.append(self._math_run(display))

    def _latex_text_to_display_text(self, text):
        display = str(text or '')
        replacements = [
            (r'\times', '×'),
            (r'\cdot', '·'),
            (r'\theta', 'θ'),
            (r'\Theta', 'θ'),
            (r'\cos', 'cos'),
            (r'\sin', 'sin'),
            (r'\tan', 'tan'),
        ]
        for source, target in replacements:
            display = display.replace(source, target)
        return display

    def _bar_math_element(self, inner):
        math = OxmlElement('m:oMath')
        bar = OxmlElement('m:bar')
        bar_pr = OxmlElement('m:barPr')
        pos = OxmlElement('m:pos')
        pos.set(qn('m:val'), 'top')
        bar_pr.append(pos)
        bar.append(bar_pr)
        expr = OxmlElement('m:e')
        expr.append(self._math_run(inner))
        bar.append(expr)
        math.append(bar)
        return math

    def _fraction_math_element(self, numerator, denominator):
        math = OxmlElement('m:oMath')
        frac = OxmlElement('m:f')
        num = OxmlElement('m:num')
        num.append(self._math_run(numerator))
        den = OxmlElement('m:den')
        den.append(self._math_run(denominator))
        frac.append(num)
        frac.append(den)
        math.append(frac)
        return math

    def _math_run(self, text):
        run = OxmlElement('m:r')
        text_element = OxmlElement('m:t')
        text_element.text = str(text or '')
        run.append(text_element)
        return run

    def _repair_parenthetical_note_artifacts(self, text):
        text = re.sub(
            (
                r'^(?P<head>.+?)\s+단\s+는\s+'
                r'(?P<desc>[^?()]+?)\s+이다\?\s*'
                r'\(\s*,\s*(?P<var>[A-Za-z][A-Za-z0-9]*)\s*'
                r'(?P<unit>\[[^\]]+\])?\s*\.\s*\)\s*$'
            ),
            self._rebuild_variable_note_before_question_mark,
            text,
        )
        text = re.sub(
            (
                r'^(?P<head>.+?\?)\s*단\s+는\s+'
                r'(?P<desc>[^()]+?)이다\s*'
                r'\(\s*,\s*(?P<var>[A-Za-z][A-Za-z0-9]*)\s*\.\s*\)\s*$'
            ),
            self._rebuild_variable_note_after_question_mark,
            text,
        )
        text = re.sub(
            (
                r'^(?P<head>.+?)\s+단\s+의\s+값(?P<angle>\d+)\s*\?\s*'
                r'\(\s*,\s*(?P<expr>sin\s*(?P=angle)˚)\s*'
                r'은\s+이다(?P<value>[0-9.]+)\s*\.\s*\)\s*$'
            ),
            self._rebuild_trig_value_note,
            text,
        )
        return text

    def _rebuild_variable_note_before_question_mark(self, match):
        desc = re.sub(r'\s+', ' ', match.group('desc')).strip()
        unit = match.group('unit') or ''
        return f"{match.group('head').strip()}? (단, {match.group('var')}는 {desc}{unit}이다.)"

    def _rebuild_variable_note_after_question_mark(self, match):
        desc = re.sub(r'\s+', ' ', match.group('desc')).strip()
        return f"{match.group('head').strip()} (단, {match.group('var')}는 {desc}이다.)"

    def _rebuild_trig_value_note(self, match):
        expr = re.sub(r'\s+', '', match.group('expr'))
        return f"{match.group('head').strip()}? (단, {expr}의 값은 {match.group('value')}이다.)"

    def _valid_spans(self, text, spans):
        normalized = []
        for span in spans:
            try:
                start = int(span.get('start'))
                end = int(span.get('end'))
            except (TypeError, ValueError, AttributeError):
                continue
            if start < 0 or end <= start or end > len(text):
                continue
            normalized.append((start, end, span))

        normalized.sort(key=lambda item: (item[0], item[1]))
        merged = []
        last_end = -1
        for start, end, span in normalized:
            if start < last_end:
                continue
            merged.append((start, end, span))
            last_end = end
        return merged

    def _add_format_tables(self, doc, format_json):
        spec = parse_format_payload(format_json)
        for table_spec in spec.get('tables') or []:
            self._add_one_format_table(doc, table_spec)

    def _add_one_format_table(self, doc, table_spec):
        table_id = table_spec.get('id', 'table')
        try:
            layout = resolve_table_layout(table_spec)
        except Exception as exc:
            logger.warning("Table layout calculation failed for %s: %s", table_id, exc)
            layout = fallback_table_layout(table_spec)
        if layout.fallback_used:
            warning = f"{table_id}:table_layout_fallback"
            if warning not in self.warnings:
                self.warnings.append(warning)
        use_wide_section = (
            self._table_requires_one_column(table_spec, layout)
            and hasattr(doc, 'add_section')
        )
        if use_wide_section:
            wide_section = doc.add_section(WD_SECTION.CONTINUOUS)
            self._set_columns(wide_section, 1)
        try:
            return self._render_one_format_table(doc, table_spec, layout)
        finally:
            if use_wide_section:
                restored_section = doc.add_section(WD_SECTION.CONTINUOUS)
                self._set_columns(restored_section, 2)

    def _render_one_format_table(self, doc, table_spec, layout):
        mode = effective_table_render_mode(table_spec, self.table_render_mode)
        if mode == 'image':
            if self._add_table_image(doc, table_spec, layout):
                return True
            self.warnings.append(
                f"{table_spec.get('id', 'table')}:image_to_native"
            )
            if self._add_native_table(doc, table_spec, layout):
                return True
        else:
            try:
                if self._add_native_table(doc, table_spec, layout):
                    return True
            except Exception as exc:
                logger.warning("Native table render failed: %s", exc)
            self.warnings.append(
                f"{table_spec.get('id', 'table')}:native_to_image"
            )
            if self._add_table_image(doc, table_spec, layout):
                return True
        self.warnings.append(f"{table_spec.get('id', 'table')}:table_unrendered")
        return False

    @staticmethod
    def _table_requires_one_column(table_spec, layout=None):
        if layout is not None:
            return bool(layout.wide)
        return bool((table_spec.get('layout') or {}).get('wide'))

    def _add_text_with_format_tables(
        self,
        doc,
        text,
        format_json=None,
        prefix='',
        size_pt=11,
        highlight=False,
    ):
        """Render text and tables in anchor order, returning the last text paragraph."""
        text = str(text or '')
        payload = parse_format_payload(format_json)
        anchored = []
        for index, table in enumerate(payload.get('tables') or []):
            offset, _recovered = resolve_table_anchor(text, table.get('anchor'))
            anchored.append((offset, index, table))
        anchored.sort(key=lambda item: (item[0], item[1]))
        if not anchored:
            paragraph = doc.add_paragraph()
            self._format_paragraph(paragraph)
            self._add_formatted_text(
                paragraph,
                text,
                format_json,
                prefix=prefix,
                size_pt=size_pt,
                highlight=highlight,
            )
            return paragraph

        cursor = 0
        last_paragraph = None
        prefix_value = prefix
        for offset, _index, table in anchored:
            offset = min(len(text), max(cursor, offset))
            if offset > cursor or prefix_value:
                last_paragraph = doc.add_paragraph()
                self._format_paragraph(last_paragraph)
                self._add_formatted_text(
                    last_paragraph,
                    text[cursor:offset],
                    self._slice_format_json(payload, cursor, offset),
                    prefix=prefix_value,
                    size_pt=size_pt,
                    highlight=highlight,
                )
                prefix_value = ''
            self._add_one_format_table(doc, table)
            cursor = offset
        if cursor < len(text):
            last_paragraph = doc.add_paragraph()
            self._format_paragraph(last_paragraph)
            self._add_formatted_text(
                last_paragraph,
                text[cursor:],
                self._slice_format_json(payload, cursor, len(text)),
                prefix=prefix_value,
                size_pt=size_pt,
                highlight=highlight,
            )
        if last_paragraph is None:
            last_paragraph = doc.add_paragraph()
            self._format_paragraph(last_paragraph)
        return last_paragraph

    @staticmethod
    def _slice_format_json(payload, start, end):
        spans = []
        for span in payload.get('spans') or []:
            try:
                span_start = int(span.get('start'))
                span_end = int(span.get('end'))
            except (TypeError, ValueError, AttributeError):
                continue
            clipped_start = max(start, span_start)
            clipped_end = min(end, span_end)
            if clipped_end <= clipped_start:
                continue
            adjusted = dict(span)
            adjusted['start'] = clipped_start - start
            adjusted['end'] = clipped_end - start
            spans.append(adjusted)
        return json.dumps({'spans': spans}, ensure_ascii=False) if spans else None

    @staticmethod
    def _width_twips(width_mm):
        return max(1, round(float(width_mm) / 25.4 * 1440))

    @classmethod
    def _set_cell_width(cls, cell, width_mm):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_width = tc_pr.find(qn('w:tcW'))
        if tc_width is None:
            tc_width = OxmlElement('w:tcW')
            tc_pr.append(tc_width)
        tc_width.set(qn('w:type'), 'dxa')
        tc_width.set(qn('w:w'), str(cls._width_twips(width_mm)))

    @classmethod
    def _set_fixed_table_widths(cls, table, widths_mm):
        table.autofit = False
        table_pr = table._tbl.tblPr
        table_layout = table_pr.find(qn('w:tblLayout'))
        if table_layout is None:
            table_layout = OxmlElement('w:tblLayout')
            table_pr.append(table_layout)
        table_layout.set(qn('w:type'), 'fixed')

        grid_columns = list(table._tbl.tblGrid)
        for index, width_mm in enumerate(widths_mm):
            if index >= len(grid_columns):
                grid_column = OxmlElement('w:gridCol')
                table._tbl.tblGrid.append(grid_column)
                grid_columns.append(grid_column)
            grid_columns[index].set(
                qn('w:w'),
                str(cls._width_twips(width_mm)),
            )
            table.columns[index].width = Mm(width_mm)

        for row in table.rows:
            for index, cell in enumerate(row.cells[:len(widths_mm)]):
                cls._set_cell_width(cell, widths_mm[index])

    def _add_native_table(self, doc, table_spec, layout=None):
        rows = [
            [str(cell or '') for cell in row]
            for row in table_spec.get('rows') or []
            if isinstance(row, list)
        ]
        if not rows:
            return False
        column_count = max((len(row) for row in rows), default=0)
        if column_count == 0:
            return False

        table = doc.add_table(rows=len(rows), cols=column_count)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        layout = layout or resolve_table_layout(table_spec)
        widths_mm = list(layout.column_widths_mm)
        if len(widths_mm) != column_count:
            layout = fallback_table_layout(table_spec)
            widths_mm = list(layout.column_widths_mm)
            warning = f"{table_spec.get('id', 'table')}:table_layout_fallback"
            if warning not in self.warnings:
                self.warnings.append(warning)
        self._set_fixed_table_widths(table, widths_mm)

        cell_specs = {
            (int(cell.get('row', -1)), int(cell.get('col', -1))): cell
            for cell in table_spec.get('cells') or []
            if isinstance(cell, dict)
        }
        merged_positions = set()
        for row_idx, row in enumerate(rows):
            for col_idx in range(column_count):
                if (row_idx, col_idx) in merged_positions:
                    continue
                cell = table.cell(row_idx, col_idx)
                spec = cell_specs.get((row_idx, col_idx), {})
                row_span = max(1, int(spec.get('row_span', 1) or 1))
                col_span = max(1, int(spec.get('col_span', 1) or 1))
                end_row = min(len(rows) - 1, row_idx + row_span - 1)
                end_col = min(column_count - 1, col_idx + col_span - 1)
                if end_row != row_idx or end_col != col_idx:
                    cell = cell.merge(table.cell(end_row, end_col))
                    for merged_row in range(row_idx, end_row + 1):
                        for merged_col in range(col_idx, end_col + 1):
                            if (merged_row, merged_col) != (row_idx, col_idx):
                                merged_positions.add((merged_row, merged_col))
                self._set_cell_width(cell, sum(widths_mm[col_idx:end_col + 1]))
                text = spec.get('text')
                if text is None:
                    text = row[col_idx] if col_idx < len(row) else ''
                paragraph = cell.paragraphs[0]
                self._format_paragraph(paragraph)
                alignment = str(spec.get('horizontal_alignment') or 'left')
                if alignment == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                vertical = str(spec.get('vertical_alignment') or 'center')
                if vertical == 'top':
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                elif vertical == 'bottom':
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                else:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                run = paragraph.add_run(str(text or ''))
                self._format_run(run, size_pt=9)
        return True

    def _add_table_image(self, doc, table_spec, layout=None):
        source = table_spec.get('source') or {}
        image_path = self._resolve_table_image_path(source.get('image_path'))
        if image_path is None:
            return False
        try:
            paragraph = doc.add_paragraph()
            self._format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = paragraph.add_run()
            width_mm = (layout or resolve_table_layout(table_spec)).total_width_mm
            run.add_picture(str(image_path), width=Mm(width_mm))
            return True
        except Exception as exc:
            logger.warning("Table source image render failed for %s: %s", image_path, exc)
            return False

    @staticmethod
    def _resolve_table_image_path(image_path):
        if not image_path:
            return None
        path = Path(str(image_path))
        candidates = [path] if path.is_absolute() else [Path.cwd() / path, path]
        return next((candidate for candidate in candidates if candidate.is_file()), None)

    def _add_image(
        self,
        doc,
        image_path,
        width_mm,
        paragraph=None,
        alignment=None,
        keep_with_next=False,
        keep_together=False,
    ):
        """Add an image paragraph when the path exists."""
        if not self._image_exists(image_path):
            return None

        try:
            if paragraph is None:
                paragraph = doc.add_paragraph()
                self._format_paragraph(paragraph, alignment=alignment)
            if keep_with_next or keep_together:
                self._keep_paragraph_together(
                    paragraph,
                    keep_with_next=keep_with_next,
                    keep_together=keep_together,
                )
            run = paragraph.add_run()
            picture_source = self._normalized_image_stream(image_path) or image_path
            run.add_picture(picture_source, width=Mm(width_mm))
            return paragraph
        except Exception as e:
            logger.error(f"Failed to add image {image_path}: {e}")
            return None

    def _image_exists(self, image_path):
        return bool(image_path and Path(image_path).exists())

    def _normalized_image_stream(self, image_path):
        """Return a PNG stream so Word export survives malformed PDF-extracted JPEG metadata."""
        try:
            from PIL import Image
        except Exception:
            return None

        try:
            with Image.open(image_path) as image:
                if image.mode not in ('RGB', 'RGBA'):
                    image = image.convert('RGB')
                else:
                    image = image.copy()

            buffer = io.BytesIO()
            image.save(buffer, format='PNG', dpi=(96, 96))
            buffer.seek(0)
            return buffer
        except Exception as exc:
            logger.debug("Image normalization failed for %s: %s", image_path, exc)
            return None

    def _keep_paragraph_together(self, paragraph, keep_with_next=False, keep_together=False):
        if keep_with_next:
            paragraph.paragraph_format.keep_with_next = True
        if keep_together:
            paragraph.paragraph_format.keep_together = True

    def _add_answer_key(self, doc, answer_key):
        """Append answer key section."""
        doc.add_page_break()

        header = doc.add_paragraph()
        self._format_paragraph(header)
        run = header.add_run("Answer Key")
        self._format_run(run, size_pt=14, bold=True)

        line = []
        for idx, ans in answer_key:
            ans_text = ans or ''
            line.append(f"{idx}. {ans_text}")
            if len(line) == 5:
                p = doc.add_paragraph()
                self._format_paragraph(p)
                run = p.add_run("  ".join(line))
                self._format_run(run, size_pt=10)
                line = []

        if line:
            p = doc.add_paragraph()
            self._format_paragraph(p)
            run = p.add_run("  ".join(line))
            self._format_run(run, size_pt=10)

    def _set_columns(self, section, cols):
        """섹션을 다단으로 설정 (OpenXML 조작)"""
        sectPr = section._sectPr
        cols_xml = sectPr.xpath('w:cols')[0]
        cols_xml.set(qn('w:num'), str(cols))
        cols_xml.set(qn('w:sep'), '1')
        cols_xml.set(qn('w:space'), self.COLUMN_SPACE_TWIPS)

    def _set_document_defaults(self, doc):
        normal = doc.styles['Normal']
        normal.font.name = self.FONT_NAME
        normal.font.size = Pt(11)
        self._set_style_font(normal, self.FONT_NAME)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = Pt(8)

    def _split_title(self, title):
        lines = [line.strip() for line in str(title).splitlines() if line.strip()]
        return lines or ['']

    def _format_paragraph(self, paragraph, alignment=None):
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(8)

        pPr = paragraph._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), self.BODY_LINE_TWIPS)
        spacing.set(qn('w:lineRule'), 'atLeast')

        rPr = pPr.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            pPr.append(rPr)
        self._set_rpr_font(rPr, self.FONT_NAME)

    def _format_run(self, run, size_pt=None, bold=False, highlight=False, underline=False):
        run.bold = bold
        if underline:
            run.underline = True
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        run.font.name = self.FONT_NAME
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        self._set_font(run, self.FONT_NAME)

    def _set_style_font(self, style, font_name):
        rPr = style._element.get_or_add_rPr()
        self._set_rpr_font(rPr, font_name)

    def _set_font(self, run, font_name):
        """한글 폰트 설정"""
        rPr = run._element.get_or_add_rPr()
        self._set_rpr_font(rPr, font_name)

    def _set_rpr_font(self, rPr, font_name):
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)

        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            rPr.append(lang)
        lang.set(qn('w:eastAsia'), 'ko-KR')
